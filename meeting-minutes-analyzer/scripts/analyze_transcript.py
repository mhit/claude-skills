# -*- coding: utf-8 -*-
import os
import re
import argparse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill, color=None, val=None):
    """
    Sets the background color of a cell.
    """
    shading_elm_1 = OxmlElement('w:shd')
    shading_elm_1.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm_1)

class PremiumDocxGenerator:
    def __init__(self, md_path, output_path):
        with open(md_path, 'r', encoding='utf-8') as f:
            self.md_content = f.read()
        self.output_path = output_path
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        # Set default font to a Japanese friendly one if possible, 
        # but stick to standard for compatibility.
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'MS Gothic'
        font.size = Pt(10.5)
        
        # Heading 1
        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'MS Gothic'
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(41, 128, 185) # Blue

        # Heading 2
        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'MS Gothic'
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(44, 62, 80) # Dark grey

    def _apply_formatting(self, paragraph, text):
        # Very basic bold parser
        parts = re.split(r'(\*\*.*\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    def generate(self):
        lines = self.md_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue

            # Title detection (bold text at top)
            if i < 2 and line.startswith('**') and line.endswith('**'):
                p = self.doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(line[2:-2])
                run.bold = True
                run.font.size = Pt(20)
                i += 1
                continue

            # Headings
            if line.startswith('# '):
                self.doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                self.doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                self.doc.add_heading(line[4:], level=3)
            
            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                p = self.doc.add_paragraph(style='List Bullet')
                self._apply_formatting(p, line[2:])
            elif re.match(r'^\d+\.', line):
                p = self.doc.add_paragraph(style='List Number')
                content = re.sub(r'^\d+\.\s*', '', line)
                self._apply_formatting(p, content)

            # Horizontal Rule
            elif line == '---':
                self.doc.add_paragraph('_' * 50)

            # Table Detection
            elif line.startswith('|'):
                # Extract table lines
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                
                if len(table_lines) >= 2:
                    self._add_table(table_lines)
                continue # i is already advanced

            # Regular text
            else:
                p = self.doc.add_paragraph()
                self._apply_formatting(p, line)
            
            i += 1

        self.doc.save(self.output_path)

    def _add_table(self, table_lines):
        # Filter out separator line (|---|---|) 
        data_rows = [l for l in table_lines if not re.match(r'^\|[:\s-]*\|', l)]
        if not data_rows: return

        # Parse cells
        rows_data = []
        for row_str in data_rows:
            # Standard markdown table parsing
            parts = row_str.strip().split('|')
            # If line starts and ends with |, the first and last split items are empty
            if len(parts) > 1 and parts[0].strip() == '' and parts[-1].strip() == '':
                parts = parts[1:-1]
            
            # Clean cells
            cells = [c.strip() for c in parts]
            rows_data.append(cells)

        if not rows_data: return
        
        num_cols = max(len(r) for r in rows_data)
        if num_cols == 0: return

        table = self.doc.add_table(rows=len(rows_data), cols=num_cols)
        table.style = 'Table Grid'

        for row_idx, row_cells_data in enumerate(rows_data):
            for col_idx, cell_text in enumerate(row_cells_data):
                if col_idx < num_cols:
                    cell = table.cell(row_idx, col_idx)
                    # Support <br> for newlines in cells
                    cell_text = cell_text.replace('<br>', '\n')
                    # Apply basic formatting
                    p = cell.paragraphs[0]
                    self._apply_formatting(p, cell_text)
                    
                    # Header styling
                    if row_idx == 0:
                        set_cell_background(cell, "BDD7EE") # Light blue
                        if p.runs:
                            p.runs[0].bold = True

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--md', required=True)
    parser.add_argument('--out', required=True)
    args = parser.parse_args()
    
    generator = PremiumDocxGenerator(args.md, args.out)
    generator.generate()
    print(f'Generated: {args.out}')

if __name__ == '__main__':
    main()
